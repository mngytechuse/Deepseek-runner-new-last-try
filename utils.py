import os
import logging
import tempfile
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import weasyprint
from flask import render_template_string
import re

# Configure logging
logging.basicConfig(level=logging.DEBUG)

def filter_content(text):
    """
    Filter out potentially controversial content
    
    Args:
        text (str): The text content to filter
        
    Returns:
        str: The filtered text
    """
    # Basic profanity filter - can be expanded
    inappropriate_terms = [
        'obscenity', 'profanity', 'slur', 'illegal', 'hate speech',
        # Add more terms as needed
    ]
    
    filtered_text = text
    for term in inappropriate_terms:
        filtered_text = re.sub(r'\b' + re.escape(term) + r'\b', '[inappropriate]', filtered_text, flags=re.IGNORECASE)
    
    return filtered_text

def create_docx(content, output_path, subject):
    """
    Create a DOCX file from the book content
    
    Args:
        content (dict): The book content
        output_path (str): The path to save the DOCX file
        subject (str): The subject of the book
        
    Returns:
        str: The path to the created DOCX file
    """
    try:
        logging.info(f"Creating DOCX file at: {output_path}")
        
        # Ensure the directory exists
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        
        doc = Document()
        
        # Set title
        title = doc.add_heading(content.get("title", f"{subject} Educational Book"), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add subject and language info
        info_para = doc.add_paragraph()
        info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info_run = info_para.add_run(f"Subject: {content.get('subject', subject)} | Language: {content.get('language', 'English')}")
        info_run.italic = True
        
        doc.add_page_break()
        
        # Add table of contents heading
        toc_heading = doc.add_heading("Table of Contents", 1)
        toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add table of contents entries
        for entry in content.get("table_of_contents", []):
            # Add chapter title
            toc_para = doc.add_paragraph()
            toc_para.add_run(entry.get("title", "")).bold = True
            
            # Add subtopics
            for subtopic in entry.get("subtopics", []):
                subtopic_para = doc.add_paragraph()
                subtopic_para.paragraph_format.left_indent = Inches(0.5)
                subtopic_para.add_run(f"• {subtopic}")
        
        doc.add_page_break()
        
        # Add each chapter
        for chapter in content.get("chapters", []):
            # Chapter title
            chapter_title = doc.add_heading(f"Chapter {chapter.get('number', '1')}: {chapter.get('title', 'Chapter')}", 1)
            chapter_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Skip introduction, subtopics, examples, and summary
            # Focus only on questions as requested
            
            # MCQs
            doc.add_heading("Multiple Choice Questions", 2)
            for i, mcq in enumerate(chapter.get("mcqs", [])):
                mcq_para = doc.add_paragraph()
                mcq_para.add_run(f"{i+1}. {mcq.get('question', '')}").bold = True
                
                doc.add_paragraph(f"A. {mcq.get('options', ['', '', '', ''])[0]}")
                doc.add_paragraph(f"B. {mcq.get('options', ['', '', '', ''])[1]}")
                doc.add_paragraph(f"C. {mcq.get('options', ['', '', '', ''])[2]}")
                doc.add_paragraph(f"D. {mcq.get('options', ['', '', '', ''])[3]}")
                
                answer_para = doc.add_paragraph()
                answer_para.add_run(f"Correct Answer: {mcq.get('correct_answer', '')}").bold = True
                
                explanation_para = doc.add_paragraph()
                explanation_para.add_run("Explanation: ").bold = True
                explanation_para.add_run(mcq.get('explanation', ''))
                
                doc.add_paragraph()  # Add some space between questions
            
            # Very Short Q&A
            doc.add_heading("Very Short Questions and Answers", 2)
            for i, qa in enumerate(chapter.get("very_short_qa", [])):
                q_para = doc.add_paragraph()
                q_para.add_run(f"{i+1}. {qa.get('question', '')}").bold = True
                
                a_para = doc.add_paragraph()
                a_para.add_run("Answer: ").bold = True
                a_para.add_run(qa.get('answer', ''))
                
                doc.add_paragraph()  # Add some space between Q&As
            
            # Short Q&A
            doc.add_heading("Short Questions and Answers", 2)
            for i, qa in enumerate(chapter.get("short_qa", [])):
                q_para = doc.add_paragraph()
                q_para.add_run(f"{i+1}. {qa.get('question', '')}").bold = True
                
                a_para = doc.add_paragraph()
                a_para.add_run("Answer: ").bold = True
                a_para.add_run(qa.get('answer', ''))
                
                doc.add_paragraph()  # Add some space between Q&As
            
            # Long Q&A
            doc.add_heading("Long Questions and Answers", 2)
            for i, qa in enumerate(chapter.get("long_qa", [])):
                q_para = doc.add_paragraph()
                q_para.add_run(f"{i+1}. {qa.get('question', '')}").bold = True
                
                a_para = doc.add_paragraph()
                a_para.add_run("Answer: ").bold = True
                a_para.add_run(qa.get('answer', ''))
                
                doc.add_paragraph()  # Add some space between Q&As
            
            # References section removed as requested
            
            # Add page break after each chapter except the last one
            if chapter.get("number") < len(content.get("chapters", [])):
                doc.add_page_break()
        
        # Save the document
        doc.save(output_path)
        
        # Verify the file was created
        if os.path.exists(output_path):
            file_size = os.path.getsize(output_path)
            logging.info(f"DOCX file created successfully at {output_path}, size: {file_size} bytes")
            return output_path
        else:
            raise Exception(f"DOCX file not created at {output_path}")
    
    except Exception as e:
        logging.error(f"Error in create_docx: {str(e)}")
        raise

def create_pdf(content, output_path, subject):
    """
    Create a PDF file from the book content
    
    Args:
        content (dict): The book content
        output_path (str): The path to save the PDF file
        subject (str): The subject of the book
        
    Returns:
        str: The path to the created PDF file
    """
    try:
        logging.info(f"Creating PDF file at: {output_path}")
        
        # Ensure the directory exists
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        
        # Create HTML template for the PDF
        html_template = """
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <title>{{ title }}</title>
            <style>
                body { 
                    font-family: Arial, sans-serif; 
                    margin: 40px;
                    line-height: 1.5;
                }
                h1 { 
                    color: #2c3e50; 
                    text-align: center;
                    font-size: 24px;
                    margin-top: 20px;
                    margin-bottom: 10px;
                }
                h2 { 
                    color: #3498db; 
                    font-size: 20px;
                    margin-top: 15px;
                    margin-bottom: 10px;
                }
                h3 { 
                    color: #2980b9; 
                    font-size: 18px;
                    margin-top: 12px;
                    margin-bottom: 8px;
                }
                p { margin: 10px 0; }
                .info {
                    text-align: center;
                    font-style: italic;
                    margin-bottom: 20px;
                }
                .toc-entry {
                    margin: 5px 0;
                    font-weight: bold;
                }
                .toc-subtopic {
                    margin: 3px 0 3px 30px;
                }
                .mcq-question {
                    font-weight: bold;
                    margin-top: 10px;
                }
                .mcq-option {
                    margin-left: 20px;
                }
                .mcq-answer {
                    font-weight: bold;
                    margin-top: 5px;
                }
                .mcq-explanation {
                    margin-top: 5px;
                    margin-bottom: 15px;
                    font-style: italic;
                }
                .qa-question {
                    font-weight: bold;
                    margin-top: 10px;
                }
                .qa-answer {
                    margin-top: 5px;
                    margin-bottom: 15px;
                }
                .reference {
                    font-size: 0.9em;
                    margin: 5px 0;
                }
                .page-break {
                    page-break-after: always;
                }
            </style>
        </head>
        <body>
            <h1>{{ title }}</h1>
            <p class="info">Subject: {{ subject }} | Language: {{ language }}</p>
            
            <div class="page-break"></div>
            
            <h1>Table of Contents</h1>
            {% for entry in table_of_contents %}
                <p class="toc-entry">{{ entry.title }}</p>
                {% for subtopic in entry.subtopics %}
                    <p class="toc-subtopic">• {{ subtopic }}</p>
                {% endfor %}
            {% endfor %}
            
            <div class="page-break"></div>
            
            {% for chapter in chapters %}
                <h1>Chapter {{ chapter.number }}: {{ chapter.title }}</h1>
                
                <!-- Theory content removed as requested -->
                
                <h2>Multiple Choice Questions</h2>
                {% for mcq in chapter.mcqs %}
                    <p class="mcq-question">{{ loop.index }}. {{ mcq.question }}</p>
                    <p class="mcq-option">A. {{ mcq.options[0] }}</p>
                    <p class="mcq-option">B. {{ mcq.options[1] }}</p>
                    <p class="mcq-option">C. {{ mcq.options[2] }}</p>
                    <p class="mcq-option">D. {{ mcq.options[3] }}</p>
                    <p class="mcq-answer">Correct Answer: {{ mcq.correct_answer }}</p>
                    <p class="mcq-explanation">Explanation: {{ mcq.explanation }}</p>
                {% endfor %}
                
                <h2>Very Short Questions and Answers</h2>
                {% for qa in chapter.very_short_qa %}
                    <p class="qa-question">{{ loop.index }}. {{ qa.question }}</p>
                    <p class="qa-answer"><strong>Answer:</strong> {{ qa.answer }}</p>
                {% endfor %}
                
                <h2>Short Questions and Answers</h2>
                {% for qa in chapter.short_qa %}
                    <p class="qa-question">{{ loop.index }}. {{ qa.question }}</p>
                    <p class="qa-answer"><strong>Answer:</strong> {{ qa.answer }}</p>
                {% endfor %}
                
                <h2>Long Questions and Answers</h2>
                {% for qa in chapter.long_qa %}
                    <p class="qa-question">{{ loop.index }}. {{ qa.question }}</p>
                    <p class="qa-answer"><strong>Answer:</strong> {{ qa.answer }}</p>
                {% endfor %}
                
                <!-- References section removed as requested -->
                
                {% if not loop.last %}
                <div class="page-break"></div>
                {% endif %}
            {% endfor %}
        </body>
        </html>
        """
        
        # Render the HTML
        html_content = render_template_string(
            html_template,
            title=content.get("title", f"{subject} Educational Book"),
            subject=content.get("subject", subject),
            language=content.get("language", "English"),
            table_of_contents=content.get("table_of_contents", []),
            chapters=content.get("chapters", [])
        )
        
        # Create a temporary HTML file
        with tempfile.NamedTemporaryFile(suffix=".html", delete=False) as temp_html:
            temp_html.write(html_content.encode('utf-8'))
            temp_html_path = temp_html.name
        
        try:
            # Convert HTML to PDF
            logging.info(f"Converting HTML to PDF using WeasyPrint")
            pdf = weasyprint.HTML(filename=temp_html_path).write_pdf()
            
            # Save the PDF
            with open(output_path, 'wb') as pdf_file:
                pdf_file.write(pdf)
            
            # Verify the file was created
            if os.path.exists(output_path):
                file_size = os.path.getsize(output_path)
                logging.info(f"PDF file created successfully at {output_path}, size: {file_size} bytes")
                return output_path
            else:
                raise Exception(f"PDF file not created at {output_path}")
        
        finally:
            # Clean up the temp HTML file in finally block to ensure it always runs
            if os.path.exists(temp_html_path):
                logging.info(f"Removing temporary HTML file: {temp_html_path}")
                os.unlink(temp_html_path)
    
    except Exception as e:
        logging.error(f"Error in create_pdf: {str(e)}")
        raise
